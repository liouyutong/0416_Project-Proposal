import docx
import re
import os

def create_docx_from_md(md_file_path, docx_file_path):
    doc = docx.Document()
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip('\n')
        
        # Heading 1
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        # Heading 2
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        # Heading 3
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        # Lists
        elif line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        # Empty line
        elif not line.strip():
            continue
        # Table (simplified: ignore for basic txt conversion, but we can do simple rows)
        elif '|' in line and '--' not in line:
            # We skip complex table formatting for now to keep the script robust
            doc.add_paragraph(line)
        # Normal paragraph
        else:
            # Basic bold parsing **text**
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

    doc.save(docx_file_path)

if __name__ == "__main__":
    md_file = r'c:\Users\user\Desktop\DIC6\產品規格書.md'
    docx_file = r'c:\Users\user\Desktop\DIC6\產品規格書.docx'
    create_docx_from_md(md_file, docx_file)
    print(f"Successfully generated {docx_file}")
